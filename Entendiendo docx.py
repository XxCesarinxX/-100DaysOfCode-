from docx import Document
from docx.shared import Inches

document = Document() 

document.add_heading('Document Title', 0)  #Titulo


p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True #negritas
p.add_run(' and some ')
p.add_run('italic.').italic = True #curveado de letra


document.add_heading('Heading, level 1', level=1)
document.add_heading('Heading, level 2', level=2)
document.add_heading('Heading, level 3', level=3)
document.add_heading('Heading, level 4', level=4)
document.add_heading('Heading, level 5', level=5)

document.add_paragraph('Intense quote', style='Intense Quote') #true=bold y italic y azul

document.add_paragraph(
    'first item in unordered list', style='List Bullet' #Listado de bbiñetas
)
document.add_paragraph(
    'first item in ordered list', style='List Number' #Listado de numeros
)

document.add_picture(r'C:\ Users \ cesar \ Pictures \ Phothoshop \ banner', width=Inches(0.25))

'''records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc'''

document.add_page_break()

document.save(r'C:\Users\cesar\Documents\GitHub\-100DaysOfCode-\ demo.docx')